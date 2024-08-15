# import streamlit as st
# import pdfplumber
# from docx import Document
# from sentence_transformers import SentenceTransformer

# from pinecone import Pinecone, Index, ServerlessSpec
# import os
# from dotenv import load_dotenv
# import openai
# # from openai import OpenAI
# # client = OpenAI()

# openai.api_key = os.getenv('OPENAI_API_KEY')

# load_dotenv()

# # Function to convert resume to text
# def convert_resume_to_text(file):
#     text = ""
#     if file.type == "application/pdf":
#         with pdfplumber.open(file) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() + "\n"
#     elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         doc = Document(file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     else:
#         st.error("Unsupported file type. Please upload a PDF or DOCX file.")
#     return text

# # Streamlit UI
# st.title("Resume Analyzer")

# # Upload resume
# resume_file = st.file_uploader("Upload Your Resume (PDF or DOCX)", type=["pdf", "docx"])

# if resume_file is not None:
#     resume_text = convert_resume_to_text(resume_file)
    
#     # Role interested in
#     role = st.text_input("Role Interested In")

#     # Job description (optional)
#     job_description = st.text_area("Job Description (Optional)")

#     # Time commitment
#     time_commitment = st.slider("How much time can you commit per week? (hours)", min_value=1, max_value=40, value=10)

#     if st.button("Submit"):
#         # Generate user vector
#         model = SentenceTransformer('all-MiniLM-L6-v2')
#         user_input = f"{role} {job_description} {resume_text} {time_commitment} hours/week"
#         user_vector = model.encode(user_input).tolist()

#         # Initialize Pinecone
#         PINECONE_API_KEY = os.getenv('PINECONE_API_KEY')
#         PINECONE_INDEX_NAME = 'interview-prep'

#         pc = Pinecone(api_key=PINECONE_API_KEY)
#         index_list = pc.list_indexes()

#         # if PINECONE_INDEX_NAME not in index_list:
#         #     pc.create_index(
#         #         name=PINECONE_INDEX_NAME,
#         #         dimension=384,
#         #         metric='cosine',
#         #         spec=ServerlessSpec(
#         #             cloud='aws',
#         #             region='us-east-1'
#         #         )
#         #     )

#         index = pc.Index(PINECONE_INDEX_NAME)

#         result = index.query(
#         vector=user_vector,
#         top_k=5,
#         include_metadata=True
#     )

#         top_courses = result['matches']

#         # Function to generate course summaries
#         def generate_course_summary(course, time_commitment):
#             # Safely access metadata or provide a default value if it's missing
#             course_title = course.get('metadata', {}).get('title', 'N/A')
#             course_url = course.get('metadata', {}).get('url', 'N/A')
#             course_description = course.get('metadata', {}).get('description', 'N/A')
#             course_duration = course.get('metadata', {}).get('duration', 'N/A')
#             course_rating = course.get('metadata', {}).get('rating', 'N/A')
#             course_viewers = course.get('metadata', {}).get('viewers', 'N/A')
            
#             prompt = f"""
#                     The user is interested in a role as '{role}' and has provided the following job description: '{job_description}'. They can commit {time_commitment} hours per week to learning. Based on this information, please provide a detailed summary of the following course:
                    
#                     - **Course Title**: {course_title}
#                     - **URL**: {course_url}
#                     - **Description**: {course_description}
#                     - **Duration**: {course_duration} hours
#                     - **Rating**: {course_rating} stars
#                     - **Viewers**: {course_viewers} people

#                     Please calculate how many weeks it would take for the user to complete this course based on their time commitment, and whether this fits the user's schedule. Additionally, summarize what the course is about, the key skills it covers, and provide any other relevant insights.
#                     Please do not give hypothetical courses that do not have a url. The structure of the output should be:
#                     Course Title: 
#                     URL: 
#                     Description:
#                     Duration:
#                     Rating:
#                     Viewers:
#                     Weekly Commitment Calculation:
#                     Total Time to complete the course:

#                     """

#             response = openai.ChatCompletion.create(
#                 model="gpt-4o",  # or 'gpt-3.5-turbo'
#                 messages=[
#                     {"role": "system", "content": "You are an assistant who provides course recommendations."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 max_tokens=300  # Adjust token limit as needed
#             )
            
#             return response['choices'][0]['message']['content'].strip()
        
            

#         # Generate summaries for the top 5 courses
#         summaries = [generate_course_summary(course, time_commitment) for course in top_courses]

        

#         # Display the top 5 courses
#         st.write("### Top 5 Recommended Courses")
#         for i, summary in enumerate(summaries):
#             st.write(f"#### Course {i + 1}")
#             st.write(summary)

        

#         # Close the Pinecone connection
#         # index.close()


import streamlit as st
import pdfplumber
from docx import Document
from sentence_transformers import SentenceTransformer

from pinecone import Pinecone, Index, ServerlessSpec
import os
from dotenv import load_dotenv
import openai
# from openai import OpenAI
# client = OpenAI()

openai.api_key = os.getenv('OPENAI_API_KEY')

load_dotenv()

# Function to convert resume to text
def convert_resume_to_text(file):
    text = ""
    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        st.error("Unsupported file type. Please upload a PDF or DOCX file.")
    return text

# Streamlit UI
st.title("Custom Learning Paths for Your Next Job Interview")
st.write("Upload your resume, desired job title, job description, and your weekly time commitment (in hours) for personalized course suggestions.")

# Upload resume
resume_file = st.file_uploader("Upload Your Resume (PDF or DOCX)", type=["pdf", "docx"])

if resume_file is not None:
    resume_text = convert_resume_to_text(resume_file)
    
    # Role interested in
    role = st.text_input("Role Interested In")

    # Job description (optional)
    job_description = st.text_area("Job Description (Optional)")

    # Time commitment
    time_commitment = st.slider("How much time can you commit per week? (hours)", min_value=1, max_value=40, value=10)

    if st.button("Submit"):
        # Generate user vector
        model = SentenceTransformer('all-MiniLM-L6-v2')
        user_input = f"{role} {job_description} {resume_text} {time_commitment} hours/week"
        user_vector = model.encode(user_input).tolist()

        # Initialize Pinecone
        PINECONE_API_KEY = os.getenv('PINECONE_API_KEY')
        PINECONE_INDEX_NAME = 'interview-prep'

        pc = Pinecone(api_key=PINECONE_API_KEY)
        index_list = pc.list_indexes()

        # if PINECONE_INDEX_NAME not in index_list:
        #     pc.create_index(
        #         name=PINECONE_INDEX_NAME,
        #         dimension=384,
        #         metric='cosine',
        #         spec=ServerlessSpec(
        #             cloud='aws',
        #             region='us-east-1'
        #         )
        #     )

        index = pc.Index(PINECONE_INDEX_NAME)

        result = index.query(
        vector=user_vector,
        top_k=5,
        include_metadata=True
    )

        top_courses = result['matches']

        def estimate_token_count(text):
            """Estimate the token count of a given text."""
            return len(text.split())

        # Function to generate course summaries
        def generate_course_summary(course, time_commitment):
            # Safely access metadata or provide a default value if it's missing
            course_title = course.get('metadata', {}).get('title', 'N/A')
            course_url = course.get('metadata', {}).get('url', 'N/A')
            course_description = course.get('metadata', {}).get('description', 'N/A')
            course_duration = course.get('metadata', {}).get('duration', 'N/A')
            course_rating = course.get('metadata', {}).get('rating', 'N/A')
            course_viewers = course.get('metadata', {}).get('viewers', 'N/A')
            
            prompt = f"""
                    The user is interested in a role as '{role}' and has provided the following job description: '{job_description}'. They can commit {time_commitment} hours per week to learning. Based on this information, please provide a detailed summary of the following course:
                    
                    - **Course Title**: {course_title}
                    - **URL**: {course_url}
                    - **Description**: {course_description}
                    - **Duration**: {course_duration} hours
                    - **Rating**: {course_rating} stars
                    - **Viewers**: {course_viewers} people

                    Please calculate how many weeks it would take for the user to complete this course based on their time commitment, and whether this fits the user's schedule. 
                    Please do not give hypothetical courses that do not have a url. The structure of the output should be:
                    1. Course Title: 
                    2. URL: 
                    3. Description:
                    4. Duration:
                    5. Rating:
                    6. Viewers:
                    7. Weekly Commitment Calculation:
                    8. Total Time to complete the course:
                    Keep the output under 250 words.
                    """

            response = openai.ChatCompletion.create(
                model="gpt-4o",  # or 'gpt-3.5-turbo'
                messages=[
                    {"role": "system", "content": "You are an assistant who provides course recommendations."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=4096  # Adjust token limit as needed
            )
            
            summary = response['choices'][0]['message']['content'].strip()

    # Check for incomplete responses or cases where AI asks for more details
            if "provide the course details" in summary or "hypothetical" in summary:
                return None

            return summary
        
            

        # Generate summaries for the top 5 courses
        summaries = []
        for course in top_courses:
            summary = generate_course_summary(course, time_commitment)
            if summary:
                summaries.append(summary)

        # Filter out any None results (where courses didn't have sufficient metadata or summary was incomplete)
        summaries = [summary for summary in summaries if summary]

        # Display the top 5 courses
        if summaries:
            st.write("### Top 5 Recommended Courses")
            for i, summary in enumerate(summaries):
                with st.expander(f"Course {i + 1}"):
                    st.write(summary)

        

        # Close the Pinecone connection
        # index.close()